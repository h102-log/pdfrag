"""PoC sample generator: build a DOCX and a PDF with text + a table.

Run once to populate samples/ for loader PoC. Same pump-manual domain as the MD.
"""
from pathlib import Path

SAMPLES = Path(__file__).resolve().parent.parent / "samples"
SAMPLES.mkdir(exist_ok=True)


def make_docx() -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("급수 펌프 P-202 점검 보고서", level=1)
    doc.add_heading("1. 설비 구성", level=2)
    doc.add_paragraph(
        "급수 펌프 P-202는 모터 M-30으로 구동되며 베어링 B-9로 축을 지지한다. "
        "P-202의 토출 라인은 보일러 BLR-1로 연결된다."
    )
    doc.add_heading("2. 점검 결과", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "항목", "측정값", "판정"
    for item, val, verdict in [
        ("베어링 B-9 진동", "4.2 mm/s", "정상"),
        ("모터 M-30 전류", "18.5 A", "정상"),
        ("실 S-8 누설", "감지됨", "주의"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = item, val, verdict
    doc.add_heading("3. 조치", level=2)
    doc.add_paragraph(
        "실 S-8 누설이 감지되어 운전팀이 1차 점검을 수행했다. "
        "누설 지속 시 P-202를 정지하고 실 S-8을 교체한다."
    )
    out = SAMPLES / "pump_report.docx"
    doc.save(out)
    print("wrote", out)


def make_pdf() -> None:
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    text = (
        "Vacuum Pump P-303 Maintenance Spec\n\n"
        "1. Overview\n"
        "Vacuum pump P-303 is driven by motor M-41 and supported by bearing B-12.\n"
        "The discharge of P-303 connects to condenser CD-2.\n\n"
        "2. Inspection Cycle\n"
        "Item                 Cycle      Owner\n"
        "Bearing B-12 vib.    weekly     mech team\n"
        "Motor M-41 insul.    monthly    elec team\n"
        "Seal S-11 leak       daily      ops team\n\n"
        "3. Procedure\n"
        "If bearing B-12 vibration exceeds limit, stop motor M-41 and check lubricant.\n"
        "If vibration persists, replace bearing B-12. Replacement is done by mech team.\n"
    )
    page.insert_text((72, 72), text, fontsize=11, fontname="helv")
    out = SAMPLES / "pump_spec.pdf"
    doc.save(out)
    doc.close()
    print("wrote", out)


if __name__ == "__main__":
    make_docx()
    make_pdf()
